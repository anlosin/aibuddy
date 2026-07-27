"""Word 文档插件 — 创建/读取 .docx 文件"""
import os

PLUGIN_INFO = {
    "name": "docx",
    "description": "创建 Word 文档(.docx)或读取已有文档内容。支持标题、段落、表格。",
    "version": "1.0",
}

TOOLS = [
    {
        "type": "function",
        "function": {
            "name": "create_docx",
            "description": "创建一个 Word 文档(.docx)。提供标题和段落内容，可包含多个章节。文件保存在程序同级目录下。",
            "parameters": {
                "type": "object",
                "properties": {
                    "filename": {
                        "type": "string",
                        "description": "文件名，如 report.docx"
                    },
                    "title": {
                        "type": "string",
                        "description": "文档标题"
                    },
                    "sections": {
                        "type": "array",
                        "items": {
                            "type": "object",
                            "properties": {
                                "heading": {"type": "string", "description": "章节标题"},
                                "content": {"type": "string", "description": "章节正文内容"}
                            },
                            "required": ["heading", "content"]
                        },
                        "description": "文档章节列表，每项包含 heading(标题) 和 content(内容)"
                    }
                },
                "required": ["filename", "title"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "read_docx",
            "description": "读取一个 Word 文档的内容，返回段落文本。",
            "parameters": {
                "type": "object",
                "properties": {
                    "filepath": {
                        "type": "string",
                        "description": "Word 文件路径"
                    }
                },
                "required": ["filepath"]
            }
        }
    },
]


def _safe_path(filepath):
    if os.path.isabs(filepath):
        return filepath
    return os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", filepath)


def execute(name, arguments):
    if name == "create_docx":
        return _create(arguments)
    if name == "read_docx":
        return _read(arguments)
    return f"未知工具: {name}"


def _create(args):
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return "错误: 请先安装 python-docx (pip install python-docx)"

    filename = args.get("filename", "document.docx")
    title = args.get("title", "未命名文档")
    sections = args.get("sections", [])

    doc = Document()

    # 文档标题
    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if sections:
        for sec in sections:
            heading = sec.get("heading", "")
            content = sec.get("content", "")
            if heading:
                doc.add_heading(heading, level=1)
            if content:
                p = doc.add_paragraph(content)
                p.style.font.size = Pt(11)
    else:
        doc.add_paragraph("（无内容）")

    path = _safe_path(filename)
    doc.save(path)
    return (
        f"Word 文档已创建: {filename}\n"
        f"路径: {os.path.abspath(path)}\n"
        f"章节数: {len(sections)}"
    )


def _read(args):
    try:
        from docx import Document
    except ImportError:
        return "错误: 请先安装 python-docx (pip install python-docx)"

    filepath = args.get("filepath", "")
    if not filepath:
        return "错误: 未提供文件路径"

    path = _safe_path(filepath)
    if not os.path.exists(path):
        return f"错误: 文件不存在 - {path}"

    try:
        doc = Document(path)
    except Exception as e:
        return f"无法打开文件: {e}"

    paragraphs = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            if p.style.name.startswith("Heading"):
                paragraphs.append(f"\n## {text}")
            else:
                paragraphs.append(text)

    content = "\n".join(paragraphs)
    if len(content) > 6000:
        content = content[:6000] + "\n...[内容已截断]"

    return f"[{filepath}]\n段落数: {len(doc.paragraphs)}\n\n{content}"
